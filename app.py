# app.py
from flask import Flask, render_template, request, jsonify, session
from flask_socketio import SocketIO
import sys
import os
from werkzeug.utils import secure_filename

from assistant import JarvisAssistant
from session_context import get_session_context, delete_session_context
from rag.chunking import split_text
from rag.embeddings import get_embeddings
from rag.vector_store import SessionVectorStore
from pypdf import PdfReader
import docx

# Database and Auth Imports
from models import db, User, Conversation, Message, Document
from flask_jwt_extended import (
    JWTManager, create_access_token, jwt_required, get_jwt_identity
)
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from config import JWT_SECRET_KEY, DATABASE_URL

# Check if the spaCy model is downloaded, exit if not
try:
    import spacy
    spacy.load("en_core_web_sm")
except OSError:
    print("spaCy model 'en_core_web_sm' not found.")
    print("Please run: python -m spacy download en_core_web_sm")
    sys.exit(1)

app = Flask(__name__)
# Configure Database and JWT Settings
app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['JWT_SECRET_KEY'] = JWT_SECRET_KEY
app.config['SECRET_KEY'] = JWT_SECRET_KEY  # for Flask session backing

db.init_app(app)
jwt = JWTManager(app)

# Explicitly use threading async mode for concurrent Socket.IO handler execution
socketio = SocketIO(app, async_mode="threading")

# Initialize Rate-Limiter (in-memory storage for simplicity)
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["500 per day", "100 per hour"],
    storage_uri="memory://"
)

# Auto-create tables in SQLite database
with app.app_context():
    db.create_all()

# Create a single, shared instance of assistant
jarvis = JarvisAssistant()
print("Jarvis Assistant has been initialized.")

@app.route('/')
def home():
    """Serves the main user interface."""
    return render_template('jarvis_ui.html')

# --- USER REGISTRATION & LOGIN REST ENDPOINTS ---

@app.route('/api/auth/register', methods=['POST'])
@limiter.limit("15 per minute")
def register():
    data = request.get_json() or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    
    if not username or not password:
        return jsonify({'error': 'Username and password are required'}), 400
        
    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Username is already taken'}), 400
        
    from werkzeug.security import generate_password_hash
    password_hash = generate_password_hash(password)
    user = User(username=username, password_hash=password_hash)
    db.session.add(user)
    db.session.commit()
    
    # Auto-create a default starting conversation
    conv = Conversation(user_id=user.id, title="Initial Chat")
    db.session.add(conv)
    db.session.commit()
    
    return jsonify({'message': 'User registered successfully'}), 201

@app.route('/api/auth/login', methods=['POST'])
@limiter.limit("15 per minute")
def login():
    data = request.get_json() or {}
    username = data.get('username', '').strip()
    password = data.get('password', '').strip()
    
    if not username or not password:
        return jsonify({'error': 'Username and password are required'}), 400
        
    user = User.query.filter_by(username=username).first()
    if not user:
        return jsonify({'error': 'Invalid username or password'}), 401
        
    from werkzeug.security import check_password_hash
    if not check_password_hash(user.password_hash, password):
        return jsonify({'error': 'Invalid username or password'}), 401
        
    access_token = create_access_token(identity=str(user.id))
    return jsonify({
        'token': access_token,
        'user': user.to_dict()
    }), 200

@app.route('/api/auth/me', methods=['GET'])
@jwt_required()
def get_me():
    user_id = get_jwt_identity()
    user = User.query.get(user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    return jsonify({'user': user.to_dict()}), 200

# --- CONVERSATION & CHAT HISTORY REST ENDPOINTS ---

@app.route('/api/conversations', methods=['GET'])
@jwt_required()
def list_conversations():
    user_id = get_jwt_identity()
    convs = Conversation.query.filter_by(user_id=user_id).order_by(Conversation.created_at.desc()).all()
    return jsonify({'conversations': [c.to_dict() for c in convs]})

@app.route('/api/conversations', methods=['POST'])
@jwt_required()
def create_conversation():
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    title = data.get('title', 'New Conversation').strip()
    if not title:
        title = 'New Conversation'
    conv = Conversation(user_id=user_id, title=title)
    db.session.add(conv)
    db.session.commit()
    return jsonify(conv.to_dict()), 201

@app.route('/api/conversations/<int:conv_id>/messages', methods=['GET'])
@jwt_required()
def get_messages(conv_id):
    user_id = get_jwt_identity()
    conv = Conversation.query.filter_by(id=conv_id, user_id=user_id).first()
    if not conv:
        return jsonify({'error': 'Conversation not found'}), 404
    msgs = Message.query.filter_by(conversation_id=conv_id).order_by(Message.timestamp.asc()).all()
    return jsonify({'messages': [m.to_dict() for m in msgs]})

@app.route('/api/conversations/<int:conv_id>', methods=['DELETE'])
@jwt_required()
def delete_conversation(conv_id):
    user_id = get_jwt_identity()
    conv = Conversation.query.filter_by(id=conv_id, user_id=user_id).first()
    if not conv:
        return jsonify({'error': 'Conversation not found'}), 404
    db.session.delete(conv)
    db.session.commit()
    
    # Clean up associated session context
    delete_session_context(f"user_{user_id}")
    return jsonify({'message': 'Conversation deleted successfully'})

# --- ANALYTICS STATS REST ENDPOINT ---

@app.route('/api/analytics/stats', methods=['GET'])
@jwt_required()
def get_analytics_stats():
    from analytics.metrics import (
        get_intent_distribution, get_average_confidence,
        get_fallback_rate, get_average_llm_latency, get_rag_query_count
    )
    return jsonify({
        'intent_distribution': get_intent_distribution(),
        'average_confidence': get_average_confidence(),
        'fallback_rate': get_fallback_rate(),
        'average_latency': get_average_llm_latency(),
        'rag_query_count': get_rag_query_count()
    })

# --- SOCKET.IO EVENT HANDLERS ---

@socketio.on('connect')
def handle_connect(auth=None):
    """Handles a new client connection and performs JWT handshake authentication if present."""
    print(f'Client connected: {request.sid}')
    
    # Resolve token from auth structure or query parameters
    token = None
    if auth and isinstance(auth, dict):
        token = auth.get('token')
    if not token:
        token = request.args.get('token')
        
    user_id = None
    if token:
        try:
            from flask_jwt_extended import decode_token
            decoded = decode_token(token)
            user_id = decoded['sub']
        except Exception as e:
            print(f"Socket connection JWT authentication failed: {e}")
            
    if user_id:
        session['user_id'] = user_id
        ctx_key = f"user_{user_id}"
        ctx = get_session_context(ctx_key)
        
        # Determine the user's active conversation
        conv = Conversation.query.filter_by(user_id=user_id).order_by(Conversation.created_at.desc()).first()
        if not conv:
            conv = Conversation(user_id=user_id, title="Initial Chat")
            db.session.add(conv)
            db.session.commit()
        ctx['active_conversation_id'] = conv.id
        
        welcome_message = f"Hello, I am {jarvis.name}. Welcome back! How can I assist you today, Sir?"
    else:
        # Fallback to local session / request-based context for unauthenticated guests (like pytest runs)
        ctx_key = request.sid
        ctx = get_session_context(ctx_key)
        welcome_message = f"Hello, I am {jarvis.name}. You can ask me for the weather, news, or say 'help' to see all commands."
        
    ctx['last_source'] = 'skill'
    ctx['last_confidence'] = 1.0
    
    socketio.emit('assistant_response', {
        'message': welcome_message,
        'state': 'IDLE',
        'source': 'skill',
        'confidence': 1.0
    })

@socketio.on('disconnect')
def handle_disconnect():
    """Handles client disconnection by cleaning up context."""
    print(f'Client disconnected: {request.sid}')
    if 'user_id' in session:
        # Retain context on Redis but clear thread cache references
        delete_session_context(f"user_{session['user_id']}")
    else:
        delete_session_context(request.sid)

@socketio.on('user_command')
def handle_user_command(data):
    """Receives a command from the UI, processes it, and sends back a response."""
    command = data.get('command')
    if not command:
        return

    print(f"Received command: {command}")
    
    conv_id = data.get('conversation_id')
    
    if 'user_id' in session:
        user_id = session['user_id']
        ctx_key = f"user_{user_id}"
        ctx = get_session_context(ctx_key)
        if conv_id:
            ctx['active_conversation_id'] = int(conv_id)
    else:
        ctx_key = request.sid
        ctx = get_session_context(ctx_key)
        
    response, new_state = jarvis.process_command(command)
    source = ctx.get('last_source', 'skill')
    confidence = ctx.get('last_confidence', 1.0)
    
    print(f"Sending response: {response} (source: {source}, conf: {confidence})")
    socketio.emit('assistant_response', {
        'message': response,
        'state': new_state,
        'source': source,
        'confidence': confidence
    })

# --- FILE UPLOAD & INDEXING ENDPOINTS ---
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path):
    ext = file_path.rsplit('.', 1)[1].lower()
    if ext == 'txt':
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif ext == 'pdf':
        reader = PdfReader(file_path)
        text_list = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_list.append(t)
        return "\n".join(text_list)
    elif ext == 'docx':
        doc_obj = docx.Document(file_path)
        return "\n".join([p.text for p in doc_obj.paragraphs])
    return ""

def rebuild_session_index(session_id):
    session_docs_dir = os.path.join("documents", session_id)
    index_path = os.path.join("documents", f"index_{session_id}.faiss")
    map_path = os.path.join("documents", f"map_{session_id}.json")
    
    # Delete index files first so SessionVectorStore initializes empty
    if os.path.exists(index_path):
        try:
            os.remove(index_path)
        except Exception:
            pass
    if os.path.exists(map_path):
        try:
            os.remove(map_path)
        except Exception:
            pass
            
    if not os.path.exists(session_docs_dir):
        return
        
    remaining_files = [f for f in os.listdir(session_docs_dir) if allowed_file(f)]
    
    if not remaining_files:
        try:
            os.rmdir(session_docs_dir)
        except Exception:
            pass
        return
        
    store = SessionVectorStore(session_id)
    all_chunks = []
    
    for filename in remaining_files:
        file_path = os.path.join(session_docs_dir, filename)
        try:
            text = extract_text_from_file(file_path)
            if text.strip():
                chunks = split_text(text)
                all_chunks.extend(chunks)
        except Exception as e:
            print(f"Error reading {filename} during index rebuild: {e}")
            
    if all_chunks:
        embeddings = get_embeddings(all_chunks)
        store.add_chunks(all_chunks, embeddings)

@app.route('/upload', methods=['POST'])
def upload_file():
    from flask_jwt_extended import verify_jwt_in_request
    from flask_jwt_extended.exceptions import JWTExtendedException
    
    user_id = None
    try:
        verify_jwt_in_request(optional=True)
        user_id = get_jwt_identity()
    except JWTExtendedException:
        pass
        
    session_id = request.form.get('session_id')
    if not user_id and not session_id:
        return jsonify({'error': 'No authentication token or session ID provided'}), 400
        
    session_key = f"user_{user_id}" if user_id else session_id
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
        
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        session_docs_dir = os.path.join("documents", session_key)
        os.makedirs(session_docs_dir, exist_ok=True)
        file_path = os.path.join(session_docs_dir, filename)
        file.save(file_path)
        
        try:
            text = extract_text_from_file(file_path)
        except Exception as e:
            return jsonify({'error': f'Failed to parse file: {str(e)}'}), 500
            
        if not text.strip():
            return jsonify({'error': 'No text could be extracted from the file'}), 400
            
        # Chunk text, embed, and index
        chunks = split_text(text)
        embeddings = get_embeddings(chunks)
        
        store = SessionVectorStore(session_key)
        store.add_chunks(chunks, embeddings)
        
        # Save document metadata if authenticated
        if user_id:
            doc_record = Document.query.filter_by(user_id=user_id, filename=filename).first()
            if not doc_record:
                doc_record = Document(
                    user_id=user_id,
                    filename=filename,
                    size_bytes=os.path.getsize(file_path)
                )
                db.session.add(doc_record)
            else:
                doc_record.size_bytes = os.path.getsize(file_path)
            db.session.commit()
        
        return jsonify({
            'message': f'File {filename} successfully indexed',
            'chunks': len(chunks)
        })
    else:
        return jsonify({'error': 'Invalid file type'}), 400

@app.route('/documents', methods=['GET'])
def list_documents():
    from flask_jwt_extended import verify_jwt_in_request
    from flask_jwt_extended.exceptions import JWTExtendedException
    
    user_id = None
    try:
        verify_jwt_in_request(optional=True)
        user_id = get_jwt_identity()
    except JWTExtendedException:
        pass
        
    session_id = request.args.get('session_id')
    if not user_id and not session_id:
        return jsonify({'error': 'No authentication token or session ID provided'}), 400
        
    session_key = f"user_{user_id}" if user_id else session_id
    
    if user_id:
        docs = Document.query.filter_by(user_id=user_id).all()
        return jsonify({'documents': [d.to_dict() for d in docs]})
    else:
        # Legacy fallback
        session_docs_dir = os.path.join("documents", session_key)
        if not os.path.exists(session_docs_dir):
            return jsonify({'documents': []})
            
        docs = []
        for filename in os.listdir(session_docs_dir):
            if allowed_file(filename):
                file_path = os.path.join(session_docs_dir, filename)
                try:
                    size = os.path.getsize(file_path)
                    docs.append({'name': filename, 'size': size})
                except Exception:
                    pass
        return jsonify({'documents': docs})

@app.route('/documents', methods=['DELETE'])
def delete_document():
    from flask_jwt_extended import verify_jwt_in_request
    from flask_jwt_extended.exceptions import JWTExtendedException
    
    user_id = None
    try:
        verify_jwt_in_request(optional=True)
        user_id = get_jwt_identity()
    except JWTExtendedException:
        pass
        
    session_id = request.args.get('session_id') or (request.json.get('session_id') if request.is_json else None)
    filename = request.args.get('filename') or (request.json.get('filename') if request.is_json else None)
    
    if not user_id and not session_id:
        return jsonify({'error': 'No authentication token or session ID provided'}), 400
    if not filename:
        return jsonify({'error': 'No filename provided'}), 400
        
    session_key = f"user_{user_id}" if user_id else session_id
    filename = secure_filename(filename)
    file_path = os.path.join("documents", session_key, filename)
    
    if user_id:
        doc_record = Document.query.filter_by(user_id=user_id, filename=filename).first()
        if not doc_record and not os.path.exists(file_path):
            return jsonify({'error': f'Document {filename} not found'}), 404
            
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
            if doc_record:
                db.session.delete(doc_record)
                db.session.commit()
            rebuild_session_index(session_key)
            return jsonify({'message': f'Document {filename} deleted and index rebuilt successfully'})
        except Exception as e:
            return jsonify({'error': f'Failed to delete document: {str(e)}'}), 500
    else:
        # Legacy fallback
        if not os.path.exists(file_path):
            return jsonify({'error': f'Document {filename} not found'}), 404
        try:
            os.remove(file_path)
            rebuild_session_index(session_key)
            return jsonify({'message': f'Document {filename} deleted and index rebuilt successfully'})
        except Exception as e:
            return jsonify({'error': f'Failed to delete document: {str(e)}'}), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    is_production = os.environ.get('RENDER') == 'true'
    
    print(f"Starting Flask server on http://0.0.0.0:{port}")
    socketio.run(
        app, 
        host='0.0.0.0',
        port=port,
        debug=not is_production,
        allow_unsafe_werkzeug=True
    )